"""
Módulo de carga de documentos.
Soporta: PDF, DOCX, TXT, MD, HTML, CSV, RTF y PDFs escaneados (OCR).

Retorna UN Document por archivo con todo el texto concatenado,
para que el chunking estructural pueda operar sobre la estructura
completa del documento (no cortado por páginas).
"""

import os
from pathlib import Path

import fitz  # pymupdf
import pytesseract
from docx import Document as DocxDocument
from langchain_core.documents import Document
from pdf2image import convert_from_path
from rich.console import Console

import config

console = Console()


class DocumentLoader:
    """Cargador universal de documentos con soporte OCR."""

    def __init__(self):
        if os.path.exists(config.TESSERACT_PATH):
            pytesseract.pytesseract.tesseract_cmd = config.TESSERACT_PATH
        self.tesseract_lang = config.TESSERACT_LANG

    # ========================================
    # API pública
    # ========================================

    def load_file(self, file_path: str | Path) -> list[Document]:
        """
        Carga un archivo y retorna una lista de Documents de LangChain.
        Retorna UN solo Document con todo el texto del archivo.
        """
        file_path = Path(file_path)

        if not file_path.exists():
            console.print(f"[red]❌ Archivo no encontrado: {file_path}[/red]")
            return []

        extension = file_path.suffix.lower()

        if extension not in config.SUPPORTED_EXTENSIONS:
            console.print(
                f"[yellow]⚠️ Extensión no soportada: {extension}[/yellow]"
            )
            return []

        console.print(f"[blue]📄 Cargando: {file_path.name}[/blue]")

        try:
            match extension:
                case ".pdf":
                    return self._load_pdf(file_path)
                case ".docx" | ".doc":
                    return self._load_docx(file_path)
                case ".txt" | ".md" | ".rtf":
                    return self._load_text(file_path)
                case ".csv":
                    return self._load_text(file_path)
                case ".html" | ".htm":
                    return self._load_html(file_path)
                case _:
                    return self._load_text(file_path)
        except Exception as e:
            console.print(
                f"[red]❌ Error al cargar {file_path.name}: {e}[/red]"
            )
            return []

    def load_directory(
        self, directory: str | Path | None = None
    ) -> list[Document]:
        """Carga todos los documentos de un directorio."""
        directory = Path(directory) if directory else config.DOCUMENTS_DIR

        if not directory.exists():
            console.print(
                f"[red]❌ Directorio no encontrado: {directory}[/red]"
            )
            return []

        all_documents: list[Document] = []

        for file_path in sorted(directory.iterdir()):
            if (
                file_path.is_file()
                and file_path.suffix.lower() in config.SUPPORTED_EXTENSIONS
            ):
                docs = self.load_file(file_path)
                all_documents.extend(docs)

        console.print(
            f"[green]✅ Se cargaron {len(all_documents)} documentos "
            f"de {directory}[/green]"
        )
        return all_documents

    # ========================================
    # Carga por tipo
    # ========================================

    def _load_pdf(self, file_path: Path) -> list[Document]:
        """
        Carga un PDF completo en UN solo Document.
        Concatena todas las páginas. Si una página es escaneada, aplica OCR.
        """
        doc = fitz.open(str(file_path))
        all_pages_text: list[str] = []
        ocr_pages: list[int] = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text", sort=True).strip()

            if not text or len(text) < 50:
                console.print(
                    f"  [yellow]🔍 Página {page_num + 1}: sin texto, "
                    f"aplicando OCR...[/yellow]"
                )
                text = self._ocr_pdf_page(file_path, page_num)
                if text:
                    ocr_pages.append(page_num + 1)

            if text.strip():
                all_pages_text.append(text)

        doc.close()

        if not all_pages_text:
            console.print(
                f"  [red]❌ No se extrajo texto de {file_path.name}[/red]"
            )
            return []

        full_text = "\n\n".join(all_pages_text)

        console.print(
            f"  [green]✅ PDF: {len(all_pages_text)} páginas extraídas"
            f"{f' ({len(ocr_pages)} con OCR)' if ocr_pages else ''}[/green]"
        )

        metadata = {
            "source": str(file_path),
            "filename": file_path.name,
            "total_pages": len(all_pages_text),
            "type": "pdf",
        }
        # ChromaDB no acepta listas vacías en metadata
        if ocr_pages:
            metadata["ocr_pages"] = str(ocr_pages)

        return [Document(page_content=full_text, metadata=metadata)]

    def _ocr_pdf_page(self, file_path: Path, page_num: int) -> str:
        """Aplica OCR a una página específica de un PDF escaneado."""
        try:
            images = convert_from_path(
                str(file_path),
                first_page=page_num + 1,
                last_page=page_num + 1,
                dpi=300,
            )
            if images:
                return pytesseract.image_to_string(
                    images[0], lang=self.tesseract_lang
                ).strip()
        except Exception as e:
            console.print(
                f"  [red]❌ Error en OCR (pág {page_num + 1}): {e}[/red]"
            )
        return ""

    def _load_docx(self, file_path: Path) -> list[Document]:
        """Carga un DOCX completo en un solo Document."""
        doc = DocxDocument(str(file_path))
        parts: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = (paragraph.style.name or "").lower()

            if "heading" in style_name or "título" in style_name:
                parts.append(f"\n{text}")
            else:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        full_text = "\n".join(parts)

        if full_text.strip():
            return [
                Document(
                    page_content=full_text,
                    metadata={
                        "source": str(file_path),
                        "filename": file_path.name,
                        "type": "docx",
                    },
                )
            ]
        return []

    def _load_text(self, file_path: Path) -> list[Document]:
        """Carga un archivo de texto plano."""
        encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]

        for encoding in encodings:
            try:
                text = file_path.read_text(encoding=encoding)
                if text.strip():
                    return [
                        Document(
                            page_content=text,
                            metadata={
                                "source": str(file_path),
                                "filename": file_path.name,
                                "type": file_path.suffix.lstrip("."),
                            },
                        )
                    ]
                return []
            except UnicodeDecodeError:
                continue

        console.print(
            f"  [red]❌ No se pudo decodificar: {file_path.name}[/red]"
        )
        return []

    def _load_html(self, file_path: Path) -> list[Document]:
        """Carga un archivo HTML extrayendo solo el texto."""
        try:
            from html.parser import HTMLParser

            class _Extractor(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.result: list[str] = []
                    self._skip = False

                def handle_starttag(self, tag, attrs):
                    if tag in ("script", "style"):
                        self._skip = True

                def handle_endtag(self, tag):
                    if tag in ("script", "style"):
                        self._skip = False

                def handle_data(self, data):
                    if not self._skip and data.strip():
                        self.result.append(data.strip())

            html_content = file_path.read_text(encoding="utf-8")
            ext = _Extractor()
            ext.feed(html_content)
            text = "\n".join(ext.result)

            if text.strip():
                return [
                    Document(
                        page_content=text,
                        metadata={
                            "source": str(file_path),
                            "filename": file_path.name,
                            "type": "html",
                        },
                    )
                ]
        except Exception as e:
            console.print(f"  [red]❌ Error al cargar HTML: {e}[/red]")

        return []
